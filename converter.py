import os
import logging
import time
from fontTools.ttLib import TTFont  # For font conversions
from PIL import Image  # For image conversions
import fitz  # PyMuPDF for PDF to image
import PyPDF2  # For PDF to text
from docx import Document  # For DOCX conversions
from pdf2docx import Converter  # For PDF to DOCX
import markdown  # For Markdown to text
from odf import text, teletype  # For ODT
from odf.opendocument import load  # For ODT
from moviepy import VideoFileClip  # Corrected import for video conversions
from pydub import AudioSegment  # For audio conversions
import zipfile  # For ZIP archives
import tarfile  # For TAR/GZ archives
import py7zr  # For 7Z archives
import csv  # For CSV
import json  # For JSON
import xml.etree.ElementTree as ET  # For XML
import openpyxl  # For XLS/XLSX

# Setup logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)


def convert_file_input(input_path, output_path, input_ext, output_ext):
    """
    Convert a file from input_ext to output_ext.
    Returns True if successful, False if unsupported or failed.
    """
    logger.debug(f"Attempting conversion: {input_ext} to {output_ext} | Input: {input_path} | Output: {output_path}")

    try:
        # Fonts
        if input_ext in ['ttf', 'otf', 'woff', 'woff2'] and output_ext in ['ttf', 'otf', 'woff', 'woff2']:
            logger.debug("Converting font")
            font = TTFont(input_path)
            font.flavor = None if output_ext in ['ttf', 'otf'] else output_ext
            font.save(output_path)
            return True

        # Images
        elif input_ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp'] and output_ext in ['png', 'jpg', 'jpeg',
                                                                                                  'gif', 'bmp', 'tiff',
                                                                                                  'webp', 'pdf']:
            logger.debug("Converting image")
            img = Image.open(input_path)
            if output_ext in ['jpg', 'jpeg'] and img.mode in ['RGBA', 'P']:
                img = img.convert('RGB')
            if output_ext == 'pdf':
                img.save(output_path, "PDF", resolution=100.0)
            else:
                img.save(output_path, output_ext.upper())
            return True
        elif input_ext == 'svg' and output_ext in ['png', 'jpg', 'jpeg', 'pdf']:
            logger.warning("SVG conversion requires cairosvg, not implemented")
            return False

        # Documents
        elif input_ext == 'pdf' and output_ext in ['png', 'jpg', 'jpeg', 'docx', 'txt']:
            logger.debug(f"Converting PDF to {output_ext}")
            if output_ext in ['png', 'jpg', 'jpeg']:
                doc = fitz.open(input_path)
                page_count = len(doc)
                logger.debug(f"PDF has {page_count} pages")

                output_dir = os.path.dirname(output_path)
                base_name = os.path.splitext(os.path.basename(output_path))[0]

                if page_count == 1:
                    # Single-page PDF: Save directly as one image
                    logger.debug("Single-page PDF detected")
                    page = doc[0]
                    pix = page.get_pixmap()
                    pix.save(output_path, output_ext)
                    logger.debug(f"Saved single page to {output_path}")
                else:
                    # Multi-page PDF: Save images and zip them
                    logger.debug(f"Multi-page PDF detected with {page_count} pages")
                    temp_images = []
                    for page_num in range(page_count):
                        page = doc[page_num]
                        pix = page.get_pixmap()
                        temp_image = os.path.join(output_dir, f"{base_name}_page_{page_num + 1}.{output_ext}")
                        logger.debug(f"Saving page {page_num + 1} to {temp_image}")
                        pix.save(temp_image, output_ext)
                        if not os.path.exists(temp_image):
                            logger.error(f"Failed to save {temp_image}")
                            return False
                        temp_images.append(temp_image)

                    # Create ZIP file with .zip extension
                    zip_path = os.path.join(output_dir, f"{base_name}.zip")
                    logger.debug(f"Creating ZIP file at {zip_path}")
                    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                        for temp_image in temp_images:
                            zipf.write(temp_image, os.path.basename(temp_image))
                            logger.debug(f"Added {temp_image} to ZIP")
                            os.remove(temp_image)  # Clean up temp files
                    logger.debug(f"ZIP file created at {zip_path}")
                    # No renaming—keep it as .zip
                doc.close()
                return True
            elif output_ext == 'txt':
                with open(input_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() or ""
                    with open(output_path, 'w', encoding='utf-8') as txt_file:
                        txt_file.write(text if text.strip() else "No readable text found in PDF")
                return True
            elif output_ext == 'docx':
                logger.debug("Using pdf2docx for PDF to DOCX conversion")
                cv = Converter(input_path)
                # Check PDF rotation and adjust
                doc = fitz.open(input_path)
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    rotation = page.rotation  # Get page rotation in degrees
                    if rotation != 0:
                        logger.debug(f"Page {page_num + 1} has rotation {rotation} degrees, correcting")
                        page.set_rotation(0)  # Normalize to upright
                doc.save("temp_rotated.pdf")  # Save corrected PDF
                doc.close()
                cv.convert(output_path, pdf_path="temp_rotated.pdf")
                cv.close()
                os.remove("temp_rotated.pdf")  # Clean up temp file
                return True

        elif input_ext in ['doc', 'docx'] and output_ext in ['docx', 'pdf', 'txt', 'rtf', 'odt']:
            logger.debug(f"Converting {input_ext} to {output_ext}")
            doc = Document(input_path)  # Assumes .doc is pre-converted to .docx
            if output_ext == 'txt':
                with open(output_path, 'w', encoding='utf-8') as txt_file:
                    for para in doc.paragraphs:
                        txt_file.write(para.text + '\n')
                return True
            elif output_ext == 'docx':
                doc.save(output_path)
                return True
            elif output_ext == 'pdf':
                # Simple DOCX to PDF using PIL (limited)
                img = Image.new('RGB', (612, 792), 'white')
                from PIL import ImageDraw
                draw = ImageDraw.Draw(img)
                text = "\n".join(para.text for para in doc.paragraphs)
                draw.text((10, 10), text[:1000], fill='black')
                img.save(output_path, 'PDF')
                return True
            else:  # rtf, odt
                logger.warning(f"{input_ext} to {output_ext} requires external tools")
                return False

        elif input_ext == 'txt' and output_ext in ['pdf', 'docx', 'rtf', 'md']:
            logger.debug(f"Converting TXT to {output_ext}")
            with open(input_path, 'r', encoding='utf-8') as txt_file:
                text = txt_file.read()
            if output_ext == 'docx':
                doc = Document()
                doc.add_paragraph(text)
                doc.save(output_path)
                return True
            elif output_ext == 'md':
                with open(output_path, 'w', encoding='utf-8') as md_file:
                    md_file.write(text)
                return True
            elif output_ext == 'pdf':
                img = Image.new('RGB', (612, 792), 'white')
                from PIL import ImageDraw
                draw = ImageDraw.Draw(img)
                draw.text((10, 10), text[:1000], fill='black')
                img.save(output_path, 'PDF')
                return True
            else:  # rtf
                logger.warning("txt to rtf requires external tools")
                return False

        elif input_ext == 'rtf' and output_ext in ['pdf', 'docx', 'txt', 'odt']:
            logger.warning("RTF conversion requires rtf2xml or similar, not implemented")
            return False

        elif input_ext == 'odt' and output_ext in ['pdf', 'docx', 'txt', 'rtf']:
            logger.debug(f"Converting ODT to {output_ext}")
            doc = load(input_path)
            text_content = teletype.extractText(doc.text)
            if output_ext == 'txt':
                with open(output_path, 'w', encoding='utf-8') as txt_file:
                    txt_file.write(text_content)
                return True
            elif output_ext == 'docx':
                docx = Document()
                docx.add_paragraph(text_content)
                docx.save(output_path)
                return True
            else:  # pdf, rtf
                logger.warning(f"odt to {output_ext} requires external tools")
                return False

        elif input_ext == 'md' and output_ext in ['pdf', 'txt']:
            logger.debug(f"Converting MD to {output_ext}")
            with open(input_path, 'r', encoding='utf-8') as md_file:
                md_content = md_file.read()
            if output_ext == 'txt':
                with open(output_path, 'w', encoding='utf-8') as txt_file:
                    txt_file.write(md_content)
                return True
            elif output_ext == 'pdf':
                img = Image.new('RGB', (612, 792), 'white')
                from PIL import ImageDraw
                draw = ImageDraw.Draw(img)
                draw.text((10, 10), md_content[:1000], fill='black')
                img.save(output_path, 'PDF')
                return True

        # Audio (using pydub)
        elif input_ext in ['mp3', 'wav', 'ogg', 'flac', 'aac', 'm4a'] and output_ext in ['mp3', 'wav', 'ogg', 'flac',
                                                                                         'aac', 'm4a']:
            logger.debug(f"Converting audio {input_ext} to {output_ext} with pydub")
            audio = AudioSegment.from_file(input_path, format=input_ext)
            audio.export(output_path, format=output_ext)
            if os.path.exists(output_path):
                logger.debug(f"Audio conversion successful: {output_path}")
                return True
            else:
                logger.error(f"Output file not created: {output_path}")
                return False

        # Video (using moviepy with proper cleanup)
        elif input_ext in ['mp4', 'avi', 'mkv', 'mov', 'wmv', 'flv', 'webm'] and output_ext in ['mp4', 'avi', 'mkv',
                                                                                                'mov', 'wmv', 'flv',
                                                                                                'webm']:
            logger.debug(f"Converting video {input_ext} to {output_ext} with moviepy")
            try:
                # Load the video file
                video_clip = VideoFileClip(input_path)

                # Determine codec and settings based on output format
                if output_ext == 'mp4':
                    video_clip.write_videofile(
                        output_path,
                        codec='libx264',  # H.264 for MP4
                        audio_codec='aac',  # AAC audio for MP4
                        ffmpeg_params=['-preset', 'fast', '-crf', '23', '-threads', '4']
                    )
                elif output_ext == 'mov':
                    video_clip.write_videofile(
                        output_path,
                        codec='libx264',  # H.264 works for MOV
                        audio_codec='aac',  # AAC audio for MOV
                        ffmpeg_params=['-preset', 'fast', '-crf', '23', '-threads', '4']
                    )
                else:
                    # For other formats (avi, mkv, wmv, flv, webm), use default settings
                    video_clip.write_videofile(
                        output_path,
                        codec='libx264',  # Default video codec
                        audio_codec='mp3',  # Default audio codec for broader compatibility
                        ffmpeg_params=['-preset', 'fast', '-crf', '23', '-threads', '4']
                    )

                # Explicitly close the video clip to release file handles
                video_clip.close()

                # Wait briefly and check if output exists, with retry mechanism
                max_attempts = 5
                for attempt in range(max_attempts):
                    if os.path.exists(output_path):
                        logger.debug(f"Video conversion successful: {output_path}")
                        return True
                    else:
                        logger.warning(
                            f"Output file not yet created (attempt {attempt + 1}/{max_attempts}): {output_path}")
                        time.sleep(1)  # Wait 1 second before retrying
                logger.error(f"Output file not created after {max_attempts} attempts: {output_path}")
                return False

            except Exception as e:
                logger.error(f"Moviepy error during video conversion: {str(e)}")
                # Ensure clip is closed even on error
                if 'video_clip' in locals():
                    video_clip.close()
                return False

        # Archives
        elif input_ext in ['zip', 'tar', 'gz', '7z'] and output_ext in ['zip', 'tar', 'gz', '7z']:
            logger.debug(f"Converting archive {input_ext} to {output_ext}")
            temp_dir = 'temp_extract'
            os.makedirs(temp_dir, exist_ok=True)
            if input_ext == 'zip':
                with zipfile.ZipFile(input_path, 'r') as z:
                    z.extractall(temp_dir)
            elif input_ext == 'tar':
                with tarfile.open(input_path, 'r') as t:
                    t.extractall(temp_dir)
            elif input_ext == 'gz':
                with tarfile.open(input_path, 'r:gz') as t:
                    t.extractall(temp_dir)
            elif input_ext == '7z':
                with py7zr.SevenZipFile(input_path, 'r') as z:
                    z.extractall(temp_dir)
            # Re-archive
            if output_ext == 'zip':
                with zipfile.ZipFile(output_path, 'w') as z:
                    for root, _, files in os.walk(temp_dir):
                        for f in files:
                            z.write(os.path.join(root, f), f)
            elif output_ext == 'tar':
                with tarfile.open(output_path, 'w') as t:
                    for root, _, files in os.walk(temp_dir):
                        for f in files:
                            t.add(os.path.join(root, f), f)
            elif output_ext == 'gz':
                with tarfile.open(output_path, 'w:gz') as t:
                    for root, _, files in os.walk(temp_dir):
                        for f in files:
                            t.add(os.path.join(root, f), f)
            elif output_ext == '7z':
                with py7zr.SevenZipFile(output_path, 'w') as z:
                    for root, _, files in os.walk(temp_dir):
                        for f in files:
                            z.write(os.path.join(root, f), f)
            # Cleanup
            for root, dirs, files in os.walk(temp_dir, topdown=False):
                for f in files:
                    os.remove(os.path.join(root, f))
                for d in dirs:
                    os.rmdir(os.path.join(root, d))
            os.rmdir(temp_dir)
            return True
        elif input_ext == 'rar' and output_ext in ['zip', 'tar', 'gz', '7z']:
            logger.warning("RAR extraction requires rarfile, not implemented")
            return False

        # Others
        elif input_ext in ['csv', 'xls', 'xlsx'] and output_ext in ['csv', 'xls', 'xlsx', 'json', 'txt']:
            logger.debug(f"Converting data file {input_ext} to {output_ext}")
            if input_ext == 'csv':
                with open(input_path, 'r', encoding='utf-8') as csv_file:
                    reader = csv.DictReader(csv_file)
                    data = list(reader)
            elif input_ext in ['xls', 'xlsx']:
                wb = openpyxl.load_workbook(input_path)
                ws = wb.active
                data = [{ws.cell(1, col).value: ws.cell(row, col).value for col in range(1, ws.max_column + 1)}
                        for row in range(2, ws.max_row + 1)]
            if output_ext == 'csv':
                with open(output_path, 'w', newline='', encoding='utf-8') as csv_file:
                    writer = csv.DictWriter(csv_file, fieldnames=data[0].keys())
                    writer.writeheader()
                    writer.writerows(data)
            elif output_ext in ['xls', 'xlsx']:
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.append(list(data[0].keys()))
                for row in data:
                    ws.append(list(row.values()))
                wb.save(output_path)
            elif output_ext == 'json':
                with open(output_path, 'w', encoding='utf-8') as json_file:
                    json.dump(data, json_file)
            elif output_ext == 'txt':
                with open(output_path, 'w', encoding='utf-8') as txt_file:
                    for row in data:
                        txt_file.write(str(row) + '\n')
            return True

        elif input_ext == 'json' and output_ext in ['csv', 'txt']:
            logger.debug(f"Converting JSON to {output_ext}")
            with open(input_path, 'r', encoding='utf-8') as json_file:
                data = json.load(json_file)
            if output_ext == 'csv':
                with open(output_path, 'w', newline='', encoding='utf-8') as csv_file:
                    writer = csv.DictWriter(csv_file, fieldnames=data[0].keys())
                    writer.writeheader()
                    writer.writerows(data)
            elif output_ext == 'txt':
                with open(output_path, 'w', encoding='utf-8') as txt_file:
                    txt_file.write(str(data))
            return True

        elif input_ext == 'xml' and output_ext in ['json', 'txt']:
            logger.debug(f"Converting XML to {output_ext}")
            tree = ET.parse(input_path)
            root = tree.getroot()
            data = {root.tag: {child.tag: child.text for child in root}}
            if output_ext == 'json':
                with open(output_path, 'w', encoding='utf-8') as json_file:
                    json.dump(data, json_file)
            elif output_ext == 'txt':
                with open(output_path, 'w', encoding='utf-8') as txt_file:
                    txt_file.write(ET.tostring(root, encoding='unicode'))
            return True

        # Unsupported
        logger.warning(f"Conversion from {input_ext} to {output_ext} not supported")
        return False

    except Exception as e:
        logger.error(f"Error during conversion from {input_ext} to {output_ext}: {str(e)}")
        return False